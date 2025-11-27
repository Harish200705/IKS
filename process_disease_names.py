#!/usr/bin/env python3
"""
Extract disease names from .docx files organized in language folders.
Matches diseases across languages and generates JSON files for MongoDB.
"""

import json
import os
import re
from pathlib import Path
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    print("❌ Error: python-docx not found.")
    print("   Attempting to install with --user flag...")
    import subprocess
    import sys
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", "python-docx"])
        from docx import Document
    except Exception as e:
        print(f"   ❌ Installation failed: {e}")
        print("\n   Please install python-docx manually using one of these methods:")
        print("   1. Using pip with --user: pip install --user python-docx")
        print("   2. Using conda: conda install -c conda-forge python-docx")
        print("   3. Or use the anaconda Python: /opt/anaconda3/bin/python3 process_disease_names.py")
        sys.exit(1)

# Base paths
BASE_PATH = Path("/Users/harish/Downloads")
LANGUAGE_FOLDERS = {
    "en": "Herbal Treatment Practices - English",
    "ta": "Herbal Treatment Practices - Tamil",
    "ml": "Herbal Treatment Practices - Malayalam",
    "hi": "Herbal Treatment Practices - Hindi"
}

# Category mappings across languages
CATEGORY_MAPPINGS = {
    "PoultryBirds": {
        "en": ["Herbal Treatment Practices for Poultry Birds", "Poultry"],
        "ta": ["3. Poultry", "Poultry"],
        "ml": ["Poultry"],
        "hi": ["Poultry"]
    },
    "CowAndBuffalo": {
        "en": ["Herbal Treatment Practices for Cows and Buffaloes", "Cattle", "Cows & Buffaloes"],
        "ta": ["1. Cattle", "Cattle"],
        "ml": ["Cows & Buffaloes"],
        "hi": ["Cattle & Buffaloes"]
    },
    "SheepGoat": {
        "en": ["Herbal Treatment Practices for Sheep and Goats", "Sheep", "Goat"],
        "ta": ["2. Sheep & Goat", "Sheep & Goat"],
        "ml": ["Sheep & Goat"],
        "hi": ["Goat & sheep", "Goat & Sheep"]
    }
}


def extract_disease_name_from_file(docx_path):
    """
    Extract disease name from a .docx file.
    The disease name is typically the first heading or title.
    Returns the name EXACTLY as it appears.
    """
    try:
        doc = Document(docx_path)
        
        # Strategy 1: Check first paragraph (often the title)
        if doc.paragraphs:
            first_text = doc.paragraphs[0].text.strip()
            if first_text and len(first_text) < 200:  # Likely a title/name
                # Remove numbering like "1. " or "1) " from start
                name = re.sub(r'^\d+[\.\)]\s*', '', first_text)
                if name:
                    return name.strip()
        
        # Strategy 2: Look for first non-empty paragraph
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) < 200:
                name = re.sub(r'^\d+[\.\)]\s*', '', text)
                if name:
                    return name.strip()
        
        # Strategy 3: Extract from filename (fallback)
        filename = docx_path.stem
        # Remove numbering from filename
        name = re.sub(r'^\d+[\.\)]\s*', '', filename)
        # Remove category suffixes like "(poultry)", "(Goat)"
        name = re.sub(r'\s*\([^)]+\)\s*$', '', name)
        return name.strip()
        
    except Exception as e:
        print(f"   ⚠️  Error reading {docx_path.name}: {e}")
        # Fallback to filename
        filename = docx_path.stem
        name = re.sub(r'^\d+[\.\)]\s*', '', filename)
        name = re.sub(r'\s*\([^)]+\)\s*$', '', name)
        return name.strip()


def find_category_folder(base_path, lang_code, category):
    """Find the category folder for a given language and category."""
    lang_folder = base_path / LANGUAGE_FOLDERS[lang_code]
    
    if not lang_folder.exists():
        return None
    
    # Check each possible category name
    category_names = CATEGORY_MAPPINGS[category].get(lang_code, [])
    
    # Direct match
    for cat_name in category_names:
        cat_path = lang_folder / cat_name
        if cat_path.exists() and cat_path.is_dir():
            return cat_path
    
    # Try to find by partial match
    for item in lang_folder.iterdir():
        if item.is_dir():
            for cat_name in category_names:
                if cat_name.lower() in item.name.lower():
                    return item
    
    return None


def get_docx_files(folder_path):
    """Get all .docx files in a folder, sorted by filename."""
    if not folder_path or not folder_path.exists():
        return []
    
    docx_files = list(folder_path.glob("*.docx"))
    docx_files.sort(key=lambda x: x.name)
    return docx_files


def match_files_across_languages(category):
    """
    Match disease files across all languages for a given category.
    Returns a list of matched file groups.
    """
    matched_groups = []
    
    # Get files for each language
    lang_files = {}
    for lang_code in LANGUAGE_FOLDERS.keys():
        cat_folder = find_category_folder(BASE_PATH, lang_code, category)
        if cat_folder:
            files = get_docx_files(cat_folder)
            lang_files[lang_code] = files
            print(f"   {lang_code.upper()}: Found {len(files)} files in {cat_folder.name}")
        else:
            lang_files[lang_code] = []
            print(f"   {lang_code.upper()}: Category folder not found")
    
    # Match files by their position/index (assuming same order)
    max_files = max(len(files) for files in lang_files.values())
    
    for idx in range(max_files):
        file_group = {}
        for lang_code in LANGUAGE_FOLDERS.keys():
            files = lang_files[lang_code]
            if idx < len(files):
                file_group[lang_code] = files[idx]
        
        if file_group:  # At least one file found
            matched_groups.append(file_group)
    
    return matched_groups


def extract_disease_from_file_group(file_group, category, disease_id):
    """
    Extract disease names from a matched file group across languages.
    Returns a disease entry dictionary.
    """
    disease_entry = {
        "id": disease_id,
        "category": category,
        "names": {
            "en": "",
            "ta": "",
            "ml": "",
            "hi": ""
        }
    }
    
    # Extract disease name from each language file
    for lang_code, file_path in file_group.items():
        disease_name = extract_disease_name_from_file(file_path)
        disease_entry["names"][lang_code] = disease_name
        print(f"      {lang_code.upper()}: {disease_name}")
    
    return disease_entry


def process_category(category):
    """Process a single category and extract all diseases."""
    print(f"\n{'='*60}")
    print(f"📂 Processing Category: {category}")
    print(f"{'='*60}")
    
    # Match files across languages
    file_groups = match_files_across_languages(category)
    
    if not file_groups:
        print(f"⚠️  No files found for category: {category}")
        return []
    
    print(f"\n   Found {len(file_groups)} disease file groups")
    print(f"   Extracting disease names...\n")
    
    diseases = []
    disease_id = 1
    
    for group_idx, file_group in enumerate(file_groups, 1):
        print(f"   Disease {group_idx}:")
        disease_entry = extract_disease_from_file_group(file_group, category, disease_id)
        diseases.append(disease_entry)
        disease_id += 1
    
    print(f"\n   ✅ Extracted {len(diseases)} diseases for {category}")
    return diseases


def main():
    """Main processing function."""
    print("🚀 Starting Disease Name Extraction")
    print(f"📁 Base path: {BASE_PATH}")
    
    # Verify base path exists
    if not BASE_PATH.exists():
        print(f"❌ Error: Base path does not exist: {BASE_PATH}")
        print("   Please update BASE_PATH in the script to point to your Downloads folder")
        return
    
    # Verify language folders exist
    missing_folders = []
    for lang_code, folder_name in LANGUAGE_FOLDERS.items():
        folder_path = BASE_PATH / folder_name
        if not folder_path.exists():
            missing_folders.append(folder_name)
    
    if missing_folders:
        print(f"⚠️  Warning: Missing language folders:")
        for folder in missing_folders:
            print(f"   - {folder}")
        print("\n   Continuing with available folders...\n")
    
    # Process each category
    categories = ["PoultryBirds", "CowAndBuffalo", "SheepGoat"]
    all_results = {}
    
    for category in categories:
        diseases = process_category(category)
        all_results[category] = diseases
    
    # Generate output files
    print(f"\n{'='*60}")
    print("📝 Generating JSON files...")
    print(f"{'='*60}\n")
    
    output_files = {
        "PoultryBirds": "poultryBirds.json",
        "CowAndBuffalo": "cowAndBuffalo.json",
        "SheepGoat": "sheepGoat.json"
    }
    
    script_dir = Path(__file__).parent
    
    for category, output_filename in output_files.items():
        if category in all_results and all_results[category]:
            output_path = script_dir / output_filename
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(all_results[category], f, ensure_ascii=False, indent=2)
            print(f"✅ Generated: {output_filename}")
            print(f"   - Location: {output_path}")
            print(f"   - Diseases: {len(all_results[category])}")
        else:
            print(f"⚠️  No data for {category}, skipping {output_filename}")
    
    print(f"\n🎉 Processing complete!")
    print(f"\n📊 Summary:")
    for category in categories:
        count = len(all_results.get(category, []))
        print(f"   - {category}: {count} diseases")
    
    return all_results


if __name__ == "__main__":
    main()
